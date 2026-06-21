"""Markdown dokümanlarını tek bir Word (.docx) belgesinde birleştirir — python-docx ile.

Sistem → Dokümanlar modülünün "Word olarak indir" özelliği kullanır. Her dosya bir Heading 1
(dosya yolu) altında; içindeki başlıklar bir kademe alta kaydırılır. Pragmatik markdown render:
başlık, paragraf (kalın/kod/italik), kod bloğu (monospace), liste, tablo. Tablolar/kod korunur.
"""
import io
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)")


def _add_inline(paragraph, text):
    """Paragrafa **kalın**, `kod`, *italik* içeren satır içi metni run'lar hâlinde ekle."""
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xB5, 0x17, 0x6C)
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def _render_markdown(doc, md):
    """Markdown metnini doc'a render et (dosya başlıkları +1 kademe kaydırılır)."""
    lines = md.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Kod bloğu (```...```)
        if stripped.startswith("```"):
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1  # kapanış ```
            for cl in (code or [""]):
                p = doc.add_paragraph()
                r = p.add_run(cl if cl else " ")
                r.font.name = "Consolas"; r.font.size = Pt(8.5)
            continue

        # Tablo (| ... |) — ardışık tablo satırlarını topla
        if stripped.startswith("|") and "|" in stripped[1:]:
            tbl_lines = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip()); i += 1
            rows = []
            for tl in tbl_lines:
                if re.match(r"^\|[\s:|-]+\|?$", tl):  # ayraç satırı (|---|---|)
                    continue
                cells = [c.strip() for c in tl.strip("|").split("|")]
                rows.append(cells)
            if rows:
                ncol = max(len(r) for r in rows)
                t = doc.add_table(rows=0, cols=ncol)
                try: t.style = "Table Grid"
                except Exception: pass
                for ri, row in enumerate(rows):
                    cells = t.add_row().cells
                    for ci in range(ncol):
                        txt = row[ci] if ci < len(row) else ""
                        cells[ci].text = ""
                        para = cells[ci].paragraphs[0]
                        _add_inline(para, txt)
                        if ri == 0:
                            for rr in para.runs: rr.bold = True
            continue

        # Başlık (#..######) — +1 kademe kaydır
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            depth = min(len(m.group(1)) + 1, 6)
            doc.add_heading(m.group(2).strip(), level=depth)
            i += 1
            continue

        # Yatay çizgi
        if re.match(r"^([-*_])\1{2,}$", stripped):
            doc.add_paragraph("─" * 30)
            i += 1
            continue

        # Liste (madde / numaralı)
        bullet = re.match(r"^\s*[-*+]\s+(.*)$", line)
        number = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if bullet:
            _add_inline(doc.add_paragraph(style="List Bullet"), bullet.group(1)); i += 1; continue
        if number:
            _add_inline(doc.add_paragraph(style="List Number"), number.group(1)); i += 1; continue

        # Boş satır
        if not stripped:
            i += 1
            continue

        # Normal paragraf
        _add_inline(doc.add_paragraph(), stripped)
        i += 1


def build_docs_docx(files):
    """files: [(relpath, markdown_içerik)] → birleşik .docx (bytes)."""
    doc = Document()
    today = datetime.now().strftime("%d.%m.%Y")
    # Kapak
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Sprenses Otel Yönetim Sistemi"); r.bold = True; r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run("Tüm Proje Dokümantasyonu"); sr.font.size = Pt(16); sr.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = d.add_run(f"{len(files)} doküman · {today}"); dr.font.size = Pt(11); dr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    for idx, (rel, content) in enumerate(files):
        doc.add_page_break()
        doc.add_heading(rel, level=1)
        try:
            _render_markdown(doc, content or "")
        except Exception as e:  # tek dosya render hatası tüm belgeyi düşürmesin
            doc.add_paragraph(f"(render hatası: {e})")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
